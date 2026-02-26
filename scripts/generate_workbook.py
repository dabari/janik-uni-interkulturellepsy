#!/usr/bin/env python3
"""
generate_workbook.py
Converts Markdown drafts (entwurf/) into a formatted Word document (ausgabe/).

Usage:
    python scripts/generate_workbook.py <Kursname>

    Kursname = Name des Unterordners in arbeiten/, z.B. InterkulturellePsy

Requirements:
    pip install python-docx>=1.1.0

Template (optional):
    Place a pre-configured Word file at scripts/template.docx to inherit
    global settings (hyphenation, document language, styles, etc.).
    If no template is found, a blank document is created instead.
"""

import json
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Path setup (repo-level constants only; course paths are set in main())
# ---------------------------------------------------------------------------
REPO_ROOT = Path(__file__).resolve().parent.parent
TEMPLATE_FILE = Path(__file__).resolve().parent / "template.docx"


# ---------------------------------------------------------------------------
# Helper: font formatting on a run
# ---------------------------------------------------------------------------
def set_run_font(run, font_size=11, bold=False, italic=False):
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rFonts.set(qn("w:cs"), "Arial")
    rpr.insert(0, rFonts)


# ---------------------------------------------------------------------------
# Helper: inline **bold** and *italic* parsing
# ---------------------------------------------------------------------------
def add_inline_formatted_runs(para, text, font_size=11):
    """Parse **bold** and *italic* markers and add styled runs to para."""
    # **bold** must be matched before *italic* to avoid partial matches
    pattern = re.compile(r'\*\*(.+?)\*\*|\*(.+?)\*')
    last_end = 0
    for m in pattern.finditer(text):
        before = text[last_end:m.start()]
        if before:
            run = para.add_run(before)
            set_run_font(run, font_size=font_size)
        if m.group(1) is not None:  # **bold**
            run = para.add_run(m.group(1))
            set_run_font(run, font_size=font_size, bold=True)
        else:  # *italic*
            run = para.add_run(m.group(2))
            set_run_font(run, font_size=font_size, italic=True)
        last_end = m.end()
    remaining = text[last_end:]
    if remaining:
        run = para.add_run(remaining)
        set_run_font(run, font_size=font_size)


# ---------------------------------------------------------------------------
# Helper: page numbers via XML field
# ---------------------------------------------------------------------------
def add_page_number(run):
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar1)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run._r.append(instrText)

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar2)


def add_footer_page_number(section, start_number=1):
    footer = section.footer
    footer.is_linked_to_previous = False  # independent footer for this section
    for para in footer.paragraphs:
        para.clear()
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    set_run_font(run, font_size=10)
    add_page_number(run)

    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:start"), str(start_number))
    pgNumType.set(qn("w:fmt"), "decimal")


# ---------------------------------------------------------------------------
# Title page
# ---------------------------------------------------------------------------
def add_title_page(doc, cfg):
    """
    Renders the title page entirely from cfg["titelblatt"].
    Each entry is either a text block or a label/value detail line:

    Text block:
      { "text": "...", "size": 12, "bold": true, "space_after": 6 }
      "{key}" placeholders in text are resolved from other config fields.

    Detail line (label: value, label bold):
      { "label": "Name", "value": "{name_vorname} {name_nachname}", "space_after": 3 }
    """
    def resolve(text):
        try:
            return text.format(**cfg)
        except KeyError:
            return text

    for item in cfg.get("titelblatt", []):
        if "label" in item:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(item.get("space_after", 3))
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            run_label = p.add_run(f"{item['label']}: ")
            set_run_font(run_label, font_size=11, bold=True)
            run_value = p.add_run(resolve(item.get("value", "")))
            set_run_font(run_value, font_size=11)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(item.get("space_after", 6))
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            run = p.add_run(resolve(item.get("text", "")))
            set_run_font(run, font_size=item.get("size", 11), bold=item.get("bold", False))


# ---------------------------------------------------------------------------
# Helper: split Aufgabenstellung from Bearbeitung
# ---------------------------------------------------------------------------
def split_aufgabenstellung(md_content):
    """
    Split markdown content at the first '---' separator.
    Expected format:
        # Aufgabe N
        Aufgabenstellung text...
        ---
        Bearbeitung text...
    Returns (heading_and_aufgabenstellung, bearbeitung_text).
    If no '---' found, everything is treated as heading + Aufgabenstellung.
    """
    lines = md_content.split("\n")
    for i, line in enumerate(lines):
        if line.strip() == "---":
            before = "\n".join(lines[:i]).strip()
            after = "\n".join(lines[i + 1:]).strip()
            return (before, after)
    return (md_content.strip(), "")


def add_aufgabenstellung(doc, text):
    """Render Aufgabenstellung text as italic Arial 11pt, justified."""
    if not text:
        return
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(line)
        set_run_font(run, font_size=11, italic=True)


# ---------------------------------------------------------------------------
# Markdown → Word (main content)
# ---------------------------------------------------------------------------
def parse_and_add_markdown(doc, md_text):
    """
    Simple Markdown renderer for main content:
    - # Heading 1  → Arial 12pt bold
    - ## Heading 2 → Arial 11pt bold
    - ### Heading 3 → Arial 11pt bold
    - **text**     → bold run
    - *text*       → italic run
    - blank line   → skip (spacing via paragraph_format)
    """
    for line in md_text.splitlines():
        line = line.rstrip()

        if line.startswith("### "):
            text = line[4:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            run = p.add_run(text)
            set_run_font(run, font_size=11, bold=True)

        elif line.startswith("## "):
            text = line[3:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            run = p.add_run(text)
            set_run_font(run, font_size=11, bold=True)

        elif line.startswith("# "):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            run = p.add_run(text)
            set_run_font(run, font_size=12, bold=True)

        elif line == "":
            pass  # paragraph spacing handles visual gaps

        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            add_inline_formatted_runs(p, line)


# ---------------------------------------------------------------------------
# Literaturverzeichnis renderer (APA 7 hanging indent)
# ---------------------------------------------------------------------------
def add_literaturverzeichnis_entries(doc, md_text):
    """
    Render bibliography entries with APA 7 hanging indent:
    first line flush left, subsequent lines indented by 1.27 cm.
    Heading lines (# ...) are skipped – the title is added separately.
    """
    for line in md_text.splitlines():
        line = line.rstrip()
        if line.startswith("#") or line == "":
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.space_after = Pt(6)
        pf.space_before = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.left_indent = Cm(1.27)
        pf.first_line_indent = Cm(-1.27)
        add_inline_formatted_runs(p, line)


# ---------------------------------------------------------------------------
# Main build function
# ---------------------------------------------------------------------------
def build_document(cfg, aufgaben, literaturverzeichnis, template_file):
    # Load template if available, otherwise start blank
    if template_file.exists():
        doc = Document(str(template_file))
        # Clear body content while preserving styles and document settings
        body = doc.element.body
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)
        print(f"Template geladen: {template_file.name}")
    else:
        doc = Document()
        print(
            "Hinweis: Keine template.docx gefunden – Dokument wird ohne Template erstellt.\n"
            f"         Lege eine Vorlage ab unter: {template_file}"
        )

    # Set Normal style base font
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    # --- Title page (section 0, no footer) ---
    # Explicitly clear section 0's footer content so no page number appears
    for para in doc.sections[0].footer.paragraphs:
        para.clear()
    add_title_page(doc, cfg)

    # --- Main content section with Arabic page numbers ---
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    add_footer_page_number(new_section, start_number=1)

    # --- Aufgaben ---
    for idx, (filename, md_content) in enumerate(aufgaben):
        if idx > 0:
            doc.add_page_break()
        header_and_aufgabe, bearbeitung = split_aufgabenstellung(md_content)
        # Render heading + Aufgabenstellung (# heading normal, rest italic)
        if header_and_aufgabe:
            ha_lines = header_and_aufgabe.split("\n")
            heading_done = False
            for line in ha_lines:
                line_s = line.rstrip()
                if not heading_done and line_s.startswith("# "):
                    # Render heading as usual (bold, 12pt)
                    text = line_s[2:].strip()
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                    run = p.add_run(text)
                    set_run_font(run, font_size=12, bold=True)
                    heading_done = True
                elif line_s == "":
                    continue
                else:
                    # Aufgabenstellung lines → italic
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                    run = p.add_run(line_s)
                    set_run_font(run, font_size=11, italic=True)
        # Render Bearbeitung as normal markdown
        if bearbeitung:
            parse_and_add_markdown(doc, bearbeitung)

    # --- Literaturverzeichnis ---
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run("Literaturverzeichnis")
    set_run_font(run, font_size=12, bold=True)

    add_literaturverzeichnis_entries(doc, literaturverzeichnis)

    return doc


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------
def main():
    if len(sys.argv) < 2:
        print("Verwendung: python scripts/generate_workbook.py <Kursname>")
        print("  Kursname = Ordnername unter arbeiten/, z.B. InterkulturellePsy")
        sys.exit(1)

    kursname = sys.argv[1]
    kurs_dir = REPO_ROOT / "arbeiten" / kursname
    entwurf_dir = kurs_dir / "entwurf"
    ausgabe_dir = kurs_dir / "ausgabe"
    config_file = entwurf_dir / "config.json"

    if not kurs_dir.exists():
        raise FileNotFoundError(f"Kursordner nicht gefunden: {kurs_dir}")
    if not config_file.exists():
        raise FileNotFoundError(f"Config nicht gefunden: {config_file}")

    with open(config_file, encoding="utf-8") as f:
        cfg = json.load(f)

    aufgabe_files = sorted(entwurf_dir.glob("aufgabe_*.md"))
    if not aufgabe_files:
        raise FileNotFoundError(f"Keine aufgabe_*.md Dateien gefunden in {entwurf_dir}")

    aufgaben = []
    for fp in aufgabe_files:
        with open(fp, encoding="utf-8") as f:
            aufgaben.append((fp.name, f.read()))
    print(f"Geladene Aufgaben: {[a[0] for a in aufgaben]}")

    lit_file = entwurf_dir / "literaturverzeichnis.md"
    if lit_file.exists():
        with open(lit_file, encoding="utf-8") as f:
            literaturverzeichnis = f.read()
    else:
        literaturverzeichnis = "_Noch keine Quellen erfasst._"

    doc = build_document(cfg, aufgaben, literaturverzeichnis, TEMPLATE_FILE)

    ausgabe_dir.mkdir(exist_ok=True)
    date_str = datetime.now().strftime("%Y%m%d")
    nachname = cfg.get("name_nachname", "Nachname")
    vorname = cfg.get("name_vorname", "Vorname")
    matnr = cfg.get("matrikelnummer", "00000000")
    kuerzel = cfg.get("kurskuerzel", "KURS")
    filename = f"{date_str}_{nachname}_{vorname}_{matnr}_{kuerzel}.docx"
    out_path = ausgabe_dir / filename

    doc.save(str(out_path))
    print(f"\nDokument gespeichert: {out_path}")
    print("\nHinweis: Öffne das Dokument in Word und aktiviere unter")
    print("         Layout → Silbentrennung → Automatisch (sofern nicht im Template gesetzt).")


if __name__ == "__main__":
    main()
